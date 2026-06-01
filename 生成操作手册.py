"""
生成用户操作手册 V2.0 (.docx)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path

doc = Document()

# ── 样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_font = h_style.font
    h_font.name = '黑体'
    h_font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1: h_font.size = Pt(16)
    elif level == 2: h_font.size = Pt(13)
    else: h_font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

def add_para(text, bold=False, size=11, align=None, color=None, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold: run.bold = True
    if color: run.font.color.rgb = color
    if align is not None: p.alignment = align
    if indent: p.paragraph_format.left_indent = Cm(0.8)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True; run.font.size = Pt(10); run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()
    return table

# ============ 封面 ============
for _ in range(6): doc.add_paragraph()
add_para('用户操作手册', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1A, 0x3C, 0x6E))
doc.add_paragraph()
add_para('基于AI规则引擎的', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('内部存货交易未实现利润自动抵消系统', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('DRP系统 T11 合并报表与一键出表 — 子模块', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
for _ in range(4): doc.add_paragraph()
add_para(f'版本：V2.0', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f'编制日期：{datetime.now().strftime("%Y年%m月%d日")}', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('小组成员：尚浩然、李金阳 · 中国社会科学院大学 MPAcc 二班', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('指导老师：张金昌', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============ 目录占位 ============
doc.add_heading('目录', level=1)
add_para('（在Word中右键此处 → 更新域 → 更新整个目录，自动生成目录）')
doc.add_page_break()

# ============ 1. 系统简介 ============
doc.add_heading('1. 系统简介', level=1)
add_para('本系统是一个基于AI规则引擎的内部存货交易未实现利润自动抵消工具，属于DRP平台T11"合并报表与一键出表"的子模块。用户只需准备三份Excel表格，系统自动完成：')
add_para('• 识别每笔内部交易的类型（顺流/逆流/平流）')
add_para('• 计算未实现内部销售利润')
add_para('• 生成标准合并抵消分录（含递延所得税 + 少数股东分摊）')
add_para('• 试算平衡验证借贷平衡')
add_para('• 支持分录手动调整与历史追溯')
add_para('• 导出抵消分录汇总报告')
add_para('适用场景：集团合并报表编制时，替代手工Excel计算内部存货交易的抵消分录。')

# ============ 2. 环境准备 ============
doc.add_heading('2. 环境准备与安装', level=1)

doc.add_heading('2.1 前置条件', level=2)
add_table(
    ['项目', '要求'],
    [
        ['操作系统', 'Windows 10/11（macOS/Linux亦可）'],
        ['Python', '3.11 或更高版本'],
        ['浏览器', 'Chrome 90+ 或 Edge 90+'],
        ['硬盘空间', '约 500MB（含依赖包）'],
    ]
)

doc.add_heading('2.2 安装步骤', level=2)
add_para('第一步：下载项目文件', bold=True)
add_para('将项目文件夹解压到本地任意目录。')
add_para('第二步：安装Python', bold=True)
add_para('打开浏览器访问 https://www.python.org/downloads/，下载并安装。关键：安装界面底部勾选"Add Python to PATH"。')
add_para('第三步：安装依赖', bold=True)
add_para('打开命令行（Win+R → 输入cmd → 回车），进入项目目录，执行：')
add_para('  pip install -r requirements.txt', indent=True)
add_para('系统自动安装 streamlit / pandas / openpyxl / reportlab，约1-2分钟。')

doc.add_heading('2.3 验证安装', level=2)
add_para('在命令行输入：streamlit --version，如果显示版本号则安装成功。')

# ============ 3. 启动系统 ============
doc.add_heading('3. 启动系统', level=1)
add_para('方式一（推荐）：双击项目目录下的"启动系统.bat"，浏览器自动打开 http://localhost:8501')
add_para('方式二：命令行进入项目目录，执行：streamlit run app.py')
add_para('若浏览器未自动打开，手动访问 http://localhost:8501。')
add_para('停止系统：关闭命令行窗口即可。')

# ============ 4. 操作流程 ============
doc.add_heading('4. 操作流程详解', level=1)

doc.add_heading('4.1 快速体验（推荐首次使用）', level=2)
add_para('① 启动系统后，点击左侧边栏"🚀 一键加载示例"按钮')
add_para('② 系统自动加载预设的股权关系、销售明细、存货结存数据')
add_para('③ 点击页面中央"🚀 执行抵消计算"按钮')
add_para('④ 查看仪表盘、图表、抵消分录等计算结果')
add_para('⑤ 可尝试下方各功能区域：试算平衡表、分录筛选、手动调整、历史记录')

doc.add_heading('4.2 步骤一：上传股权关系配置表', level=2)
add_para('目的：告诉系统集团内各公司之间的母子关系。')
add_para('点击左侧第一个上传区域，选择Excel文件。')
add_table(
    ['列名', '说明', '示例'],
    [
        ['母公司编码', '投资方主体唯一标识', 'P001'],
        ['子公司编码', '被投资方主体唯一标识', 'S001'],
        ['母公司持股比例', '0-1之间的小数', '0.80'],
        ['是否纳入合并', '是/否', '是'],
    ]
)
add_para('操作要点：这是"一次性配置"，首次上传后系统自动记忆，下次启动无需重复上传。持股比例填写小数（如80%填0.80）。')

doc.add_heading('4.3 步骤二：上传内部销售明细表', level=2)
add_para('目的：提供本月合并范围内各主体之间的存货销售数据。')
add_table(
    ['列名', '说明', '示例'],
    [
        ['销售方编码', '销售方主体标识', 'P001'],
        ['购买方编码', '购买方主体标识', 'S001'],
        ['存货批次号', '存货批次唯一标识', 'B202605001'],
        ['存货名称', '存货品名', 'A产品'],
        ['销售收入', '不含税金额（元）', '1000000'],
        ['销售成本', '不含税成本（元）', '600000'],
        ['销售数量', '件数', '100'],
        ['会计期间', 'YYYYMM格式', '202605'],
    ]
)

doc.add_heading('4.4 步骤三：上传存货期末结存表', level=2)
add_para('目的：告诉系统每批内部购入的存货还有多少没卖出去。')
add_table(
    ['列名', '说明', '示例'],
    [
        ['主体编码', '存货持有方（即购买方）', 'S001'],
        ['存货批次号', '关联销售明细表', 'B202605001'],
        ['期末结存数量', '月末尚未对外售出数量', '70'],
        ['内部购入总数量', '该批次初始购入总数', '100'],
    ]
)
add_para('注意：若某批次已全部售出（结存=0），可不填该行，系统自动将其标记为"已全部实现，无需抵消"。')

doc.add_heading('4.5 辅助：手动录入交易', level=2)
add_para('除Excel批量上传外，系统支持直接在Web界面填写单笔交易信息。')
add_para('展开"✏️ 手动录入交易（可选）"区域，填写销售方、购买方、批次号、品名、收入、成本、数量、结存后点击"添加到列表"。')
add_para('适用于：演示时临时补充数据、测试边界场景。')

doc.add_heading('4.6 执行抵消计算', level=2)
add_para('三份表格上传完成后，点击页面中央"🚀 执行抵消计算"按钮。')
add_para('系统自动完成：交易类型识别 → 未实现利润计算 → 抵消分录生成 → 哈希存证。')
add_para('计算耗时：100笔交易以内约2-5秒。')

doc.add_heading('4.7 查看仪表盘', level=2)
add_para('计算完成后，页面展示以下区域：')
add_para('① 快捷操作卡片（4个）：银行对账/关联交易/合并抵消/生成报告（前两个预留）')
add_para('② 最近操作时间线：最近8笔分录记录')
add_para('③ 异常/警告：如发现数据问题自动提示')
add_para('④ KPI指标卡片（5个）：本期交易额、未实现利润总额、已抵消笔数、冲减存货、合并净利润影响')
add_para('⑤ 抵消金额构成图：冲减存货/递延所得税资产/少数股东权益的水平柱状图')
add_para('⑥ 交易类型分布图：顺流/逆流/平流各类未实现利润占比')
add_para('⑦ 抵消前后对比表：营收/成本/存货/递延所得税/净利润五项目的抵消前后变化')

doc.add_heading('4.8 查看试算平衡表', level=2)
add_para('位于抵消前后对比表下方，自动汇总所有分录的借方合计与贷方合计。')
add_para('显示分录笔数、借方合计、贷方合计、差额。差额为0时显示"✅ 借贷平衡"（正常状态）。')
add_para('若显示"⚠ 借贷不平"，请检查是否进行了手动调整导致借贷不等。')

doc.add_heading('4.9 分录筛选与查看', level=2)
add_para('在"📝 抵消分录明细"区域，顶部有筛选标签（全部/顺流交易/逆流交易/平流交易）。')
add_para('点击标签按类型过滤分录，显示"N/M笔"。默认显示全部。')
add_para('每笔分录以卡片形式展示，包含：')
add_para('• 分录编号（ELIM-YYYYMM-NNNN格式）')
add_para('• 交易类型标签 + 交易双方 + 存货名称')
add_para('• 未实现利润金额')
add_para('• 计算参数（毛利率、未售比例、收入、成本、少数股东比例）')
add_para('• 标准借贷分录代码块')

doc.add_heading('4.10 手动调整分录', level=2)
add_para('点击分录卡片内的"✏️ 手动调整分录"展开编辑区域：')
add_para('• 修改金额：每行分录旁有金额输入框，修改后实时生效')
add_para('• 删除分录行：点击行尾"🗑️"按钮删除单行')
add_para('• 追加分录：在底部输入借方科目、贷方科目、金额后点击"添加"')
add_para('• 删除整笔：点击卡片底部"🗑️ 删除此笔抵消"删除整笔记录')
add_para('人工追加的分录标注"✏️ 人工追加"备注。注意：手动调整后建议重新核对试算平衡。')

doc.add_heading('4.11 查看历史记录', level=2)
add_para('在分录明细下方"📜 历史处理记录"区域，按批次分组展示历次计算记录。')
add_para('点击批次卡片展开，查看每笔分录的完整计算过程、分录内容和SHA256哈希值。')
add_para('哈希链机制：每笔分录哈希含上一笔哈希，形成防篡改链条。')

doc.add_heading('4.12 导出结果', level=2)
add_para('页面底部提供两种导出方式：')
add_para('• 📥 导出Excel：包含汇总表（KPI指标）+ 抵消分录明细表')
add_para('• 📥 导出PDF报告：格式化抵消分录汇总报告，含异常检测结果')
add_para('文件名自动包含会计期间，如"抵消分录明细_202605.xlsx"。')

doc.add_heading('4.13 调节税率', level=2)
add_para('在左侧边栏"⚙️ 配置"区域，拖动"所得税税率"滑块可调整税率（默认25%，高新技术企业可调为15%）。')
add_para('修改后需重新点击"执行抵消计算"。')

# ============ 5. 常见问题 ============
doc.add_heading('5. 常见问题排查', level=1)

add_para('Q1：上传Excel后提示"列名不匹配"？', bold=True)
add_para('请对照附录A的模板格式检查列名，列名必须完全一致（包括空格和标点）。可从左侧边栏下载标准模板。')

add_para('Q2：计算结果中某笔交易标记为"未识别"？', bold=True)
add_para('说明该交易的双方在股权关系表中找不到母子关联。请检查双方编码是否在股权关系表中存在。')

add_para('Q3：如何清空历史数据重新开始？', bold=True)
add_para('点击左侧边栏"🔄 清除全部"按钮，或删除项目目录下 data/ 文件夹后重启。')

add_para('Q4：税率可以修改吗？', bold=True)
add_para('可以。在左侧边栏"⚙️ 配置"中修改所得税税率（默认25%）。修改后需重新点击"执行抵消计算"。')

add_para('Q5：试算不平衡怎么办？', bold=True)
add_para('首先检查是否进行了手动调整导致借贷不等。正常由系统自动生成的分录始终借贷平衡。若手动修改导致不平衡，可通过调整分录金额或删除追加的分录来恢复。')

add_para('Q6：导出的Excel/PDF是否包含手动调整？', bold=True)
add_para('V2.0版本的导出功能反映当前页面显示的分录数据，包括手动调整后的金额和人工追加的分录。建议在导出前确认试算平衡。')

# ============ 附录 ============
doc.add_heading('附录A：Excel模板说明', level=1)
add_para('系统提供3份空白模板（可从左侧边栏下载）：')
add_table(
    ['模板文件', '用途', '频率'],
    [
        ['模板-股权关系配置表.xlsx', '集团股权结构', '一次性'],
        ['模板-内部销售明细表.xlsx', '本月内部交易数据', '每月'],
        ['模板-存货期末结存表.xlsx', '内部购入存货结存', '每月'],
    ]
)

doc.add_heading('附录B：快捷键与操作技巧', level=1)
add_para('• 一键加载示例：左侧边栏最快捷的体验方式')
add_para('• 分录筛选标签：快速定位特定类型交易')
add_para('• 卡片折叠/展开：点击分录编号展开查看详情，默认仅展开第一笔')
add_para('• 侧边栏状态：实时显示三表上传进度（0/3 → 3/3）')

# ── 保存 ──
output = Path(__file__).parent / 'T11-用户操作手册-v2.docx'
doc.save(str(output))
print(f'✅ 操作手册已生成：{output}')
