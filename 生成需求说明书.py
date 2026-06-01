"""
生成系统需求说明书 V4.0 (.docx)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_font = h_style.font
    h_font.name = '黑体'
    h_font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1: h_font.size = Pt(16)
    elif level == 2: h_font.size = Pt(13)
    else: h_font.size = Pt(11)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

def add_para(text, bold=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold: run.bold = True
    if color: run.font.color.rgb = color
    if align is not None: p.alignment = align
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()
    return table

# ============ 封面 ============
for _ in range(6): doc.add_paragraph()
add_para('软件需求规格说明书（SRS）', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1A, 0x3C, 0x6E))
doc.add_paragraph()
add_para('基于AI规则引擎的', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('内部存货交易未实现利润自动抵消系统', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('DRP系统 T11 合并报表与一键出表 — 子模块', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
for _ in range(4): doc.add_paragraph()
add_para(f'文档编号：DRP-T11-SRS-V4.0', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f'版本号：V4.0（交付定稿）', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f'编制日期：{datetime.now().strftime("%Y年%m月%d日")}', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('所属课程：财务数智化转型与AI应用（2025级会计专硕）', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('编制人：尚浩然 · 中国社会科学院大学 MPAcc 二班', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============ 版本历史 ============
doc.add_heading('版本历史', level=1)
add_table(
    ['版本', '日期', '修改内容', '作者'],
    [
        ['V1.0', '2026-05-14', '初稿：Excel上传+三种交易类型+分录生成', '尚浩然'],
        ['V2.0', '2026-05-22', '模块化架构（engine/utils）+异常检测+仪表盘+手动录入', '尚浩然'],
        ['V3.0', '2026-05-29', '补充用例图/时序图/数据模型/政策依据，交付定稿', '尚浩然'],
        ['V4.0', datetime.now().strftime('%Y-%m-%d'), '新增试算平衡表/历史记录/分录手动调整/分录筛选/驾驶舱增强', '尚浩然'],
    ]
)

# ============ 1. 引言 ============
doc.add_heading('1. 引言', level=1)

doc.add_heading('1.1 编写目的', level=2)
add_para('本文档旨在系统阐述DRP平台T11"合并报表与一键出表"下属子模块——"内部存货交易未实现利润自动抵消系统"的完整功能需求与非功能需求。本模块聚焦于合并报表编制中内部存货交易的未实现利润抵消环节，通过AI规则引擎与自动计算技术，实现从数据导入到抵消分录生成的全流程自动化。')
add_para('本文档面向以下读者群体：')
add_para('• 开发团队：作为功能实现与系统设计的唯一需求依据')
add_para('• 测试团队：依据本文档编制测试用例与验收标准')
add_para('• 项目评审方（授课教师）：评估本模块对DRP平台总体设计目标的覆盖程度')
add_para('• 财务终用户：确认系统功能与实际业务场景的匹配度')

doc.add_heading('1.2 范围声明', level=2)
add_para('本系统聚焦于T11"集团合并报表"中"合并抵消"下的"未实现利润抵消"子场景，具体覆盖：', bold=True)
add_para('✅ 内部存货交易（顺流/逆流/平流）的未实现利润自动抵消')
add_para('✅ 递延所得税自动计算')
add_para('✅ 少数股东损益分摊（逆流/平流场景）')
add_para('✅ 试算平衡验证')
add_para('✅ 抵消分录手动调整与历史追溯')
add_para('❌ 不包含：投资成本抵消、内部往来（债权债务）抵消、外币报表折算、完整合并报表生成（属于T11其他子模块或后续版本范围）')

doc.add_heading('1.3 政策依据', level=2)
add_para('• 国务院国资委《关于推动中央企业加快财务数智化转型升级的指导意见》（2026年1号文）——三（八）深化人工智能应用')
add_para('• 国务院国资委《关于加强中央企业穿透式监管的指导意见》（2026年2号文）——二（二）数据自动采集、模型自动分析、风险自动预警')
add_para('• 《企业会计准则第33号——合并财务报表》（财政部2014年修订）')
add_para('• 《企业会计准则第18号——所得税》')

doc.add_heading('1.4 术语定义', level=2)
add_table(
    ['术语', '定义'],
    [
        ['DRP', '全域数字化资源管理平台（Digital Resource Platform）'],
        ['T11', '合并报表与一键出表系统'],
        ['未实现内部销售利润', '内部销售方已确认利润，但从集团整体看存货尚未对外售出，该利润未实现'],
        ['顺流交易（DOWN）', '母公司向子公司销售'],
        ['逆流交易（UP）', '子公司向母公司销售'],
        ['平流交易（FLAT）', '子公司之间相互销售'],
        ['AI规则引擎', '基于可配置JSON规则（非硬编码）自动判断交易类型、选择抵消模板、执行计算的核心组件'],
    ]
)

# ============ 2. 项目概述 ============
doc.add_heading('2. 项目概述', level=1)

doc.add_heading('2.1 在DRP平台中的定位', level=2)
add_para('DRP系统采用六层架构（L0–L5），包含34个任务模块。本模块属于：')
add_para('L2 核心财务应用层 → 共享财务系统组（T09–T13）→ T11 合并报表与一键出表 → 内部存货交易未实现利润自动抵消子模块')
add_para('根据总体设计说明书，T11的核心定位为"对外披露"，功能涵盖：集团抵消、外币折算、合并报表生成、区块链防篡改存证。本模块聚焦其中"集团抵消"下的内部存货交易场景。')
add_para('上游依赖：T10（智慧核算引擎——获取账务数据）、T05（主数据管理——股权关系）')
add_para('下游协作：T12（电子会计档案——归档抵消分录）、T21（会计穿透监管——抵消数据查询）')
add_para('区块链依赖：T33（区块链防篡改服务——哈希上链，V1.0以本地哈希链预置扩展接口）')

doc.add_heading('2.2 产品功能概要', level=2)
add_para('本系统包含以下核心功能：')
add_para('1. 数据导入：Excel三表上传（股权关系/销售明细/存货结存）+ 手动录入')
add_para('2. 数据校验与异常检测：4种自动检测规则（库存缺失/负毛利/股权缺失/结存异常）')
add_para('3. 交易类型智能识别：基于股权关系图谱自动判定顺流/逆流/平流')
add_para('4. 未实现利润自动计算：毛利率×未售比例，支持个别计价法/FIFO')
add_para('5. 抵消分录自动生成：JSON规则引擎驱动，含递延所得税+少数股东分摊')
add_para('6. 试算平衡验证：借方/贷方合计自动对比，验证借贷平衡')
add_para('7. 分录手动调整：修改金额、删除分录行、追加入工分录')
add_para('8. 分录筛选：按交易类型（顺流/逆流/平流）过滤')
add_para('9. 可视化仪表盘：KPI指标卡+柱状图+抵消前后对比表')
add_para('10. 历史记录：按批次查看历次计算的分录详情+哈希追溯')
add_para('11. 结果导出：Excel（汇总额+明细）+ PDF报告')

doc.add_heading('2.3 用户角色', level=2)
add_table(
    ['角色', '操作', '频率'],
    [
        ['合并报表编制人', '上传Excel → 查看结果 → 调整分录 → 导出', '月度'],
        ['财务审核人', '查看抵消分录明细与试算平衡，复核调整项', '月度'],
        ['审计人员', '追溯历史批次，验证计算过程与哈希链', '按需'],
    ]
)

# ============ 3. 功能需求 ============
doc.add_heading('3. 功能需求', level=1)

doc.add_heading('3.1 数据导入（SRS-F01）', level=2)
add_para('用户通过Streamlit界面按标准模板上传三份Excel表格：')
add_table(
    ['表格', '关键列', '说明'],
    [
        ['内部销售明细表', '销售方/购买方编码、批次号、品名、收入、成本、数量、期间', '每月上传'],
        ['存货期末结存表', '主体编码、批次号、期末结存数量、内部购入总数量', '每月上传'],
        ['股权关系配置表', '母公司/子公司编码、持股比例、是否纳入合并', '一次性配置，自动记忆'],
    ]
)
add_para('辅助功能：支持在Web界面手动录入单笔交易，适用于演示时临时补充数据。')

doc.add_heading('3.2 数据校验与异常检测（SRS-F02）', level=2)
add_para('系统自动执行4种异常检测规则：')
add_table(
    ['检测规则', '触发条件', '行为'],
    [
        ['库存数据缺失', '销售明细中的批次在结存表中无记录', '⚠ 警告：可能已全部售出'],
        ['毛利率异常', '销售成本 > 销售收入', '⚠ 警告：可能为亏损销售'],
        ['股权关系缺失', '交易双方不在股权关系表中', '⚠ 警告：无法判定交易类型'],
        ['结存数据异常', '期末结存 > 购入总量', '⚠ 警告：数据逻辑矛盾'],
    ]
)
add_para('异常不阻断计算流程，按保守策略处理并汇总展示。')

doc.add_heading('3.3 交易类型智能识别（SRS-F03）', level=2)
add_para('基于股权关系图谱的最近公共祖先（LCA）判定算法：')
add_para('规则1：若销售方是购买方的母公司 → 顺流交易（DOWN）')
add_para('规则2：若购买方是销售方的母公司 → 逆流交易（UP）')
add_para('规则3：若双方有共同母公司 → 平流交易（FLAT）')
add_para('规则4：无法判定 → 标记为UNKNOWN，按顺流处理（保守策略）')

doc.add_heading('3.4 未实现利润计算（SRS-F04）', level=2)
add_para('核心公式：未实现利润 = 销售收入 × 毛利率 × 未对外出售比例')
add_para('毛利率 = (销售收入 - 销售成本) / 销售收入')
add_para('未对外出售比例 = 期末结存数量 / 内部购入总数量')
add_para('三种场景：')
add_para('• 全未售（结存=购入）：全额抵消，未实现利润 = 收入 - 成本')
add_para('• 部分售出：仅抵消未售部分')
add_para('• 全售出（结存=0）：已实现，自动跳过不抵消')
add_para('支持存货计价方法：个别计价法（按批次匹配）、先进先出法（FIFO）')

doc.add_heading('3.5 抵消分录自动生成（SRS-F05）★核心', level=2)
add_para('系统基于JSON规则引擎生成标准借贷分录。三套模板预配置于rules.json：')
add_para('【顺流交易】借：营业收入 贷：营业成本 贷：存货 | 借：递延所得税资产 贷：所得税费用')
add_para('【逆流交易】同上 + 借：少数股东权益 贷：少数股东损益（按(1-税率)×持股比例分摊）')
add_para('【平流交易】与逆流同理，按销售方少数股东持股比例分摊')
add_para('设计理念：抵消规则以JSON声明式存储，新增抵消场景只需增加配置条目，核心引擎代码无需修改。')
add_para('每笔分录自动分配编号（ELIM-YYYYMM-NNNN格式），关联计算参数（毛利率、未售比例）。')

doc.add_heading('3.6 试算平衡验证（SRS-F06）★V4.0新增', level=2)
add_para('系统在所有分录生成后自动汇总借方合计与贷方合计，验证借贷平衡。')
add_para('展示内容：分录笔数、借方合计、贷方合计、差额。差额为0显示"✅ 借贷平衡"，否则显示"⚠ 借贷不平"并标红。')
add_para('由于分录模板对称设计（每笔借必有对应贷），正常情况下差额恒为0，作为数据正确性的自我验证。')

doc.add_heading('3.7 分录筛选与手动调整（SRS-F07）★V4.0新增', level=2)
add_para('分录筛选：支持按交易类型（全部/顺流交易/逆流交易/平流交易）过滤分录卡片，实时显示"N/M笔"。')
add_para('手动调整：')
add_para('• 修改金额：每行分录可单独修改金额')
add_para('• 删除分录行：逐行删除')
add_para('• 追加分录：人工录入借贷科目和金额')
add_para('• 删除整笔：删除整笔抵消记录')
add_para('所有人工调整标注"✏️ 人工追加"备注，与系统自动生成区分。')

doc.add_heading('3.8 可视化仪表盘（SRS-F08）', level=2)
add_table(
    ['组件', '内容', '说明'],
    [
        ['KPI指标卡片', '交易额/未实现利润/抵消笔数/冲减存货/净利影响', '5个metric卡片'],
        ['快捷操作卡片', '银行对账/关联交易/合并抵消/生成报告', '4个入口（部分预留）'],
        ['最近操作时间线', '最近8条分录记录', '从calc_log表读取'],
        ['抵消金额构成图', '冲减存货/递延所得税/少数股东权益', '水平柱状图'],
        ['交易类型分布图', '顺流/逆流/平流未实现利润占比', '水平柱状图'],
        ['抵消前后对比表', '营收/成本/存货/递延税/净利润五项', 'DataFrame对比'],
        ['试算平衡表', '借贷合计+差额', '4个metric卡片'],
    ]
)

doc.add_heading('3.9 历史记录追溯（SRS-F09）★V4.0新增', level=2)
add_para('所有计算批次自动存入SQLite calc_log表。历史记录区域按批次分组展示：')
add_para('• 批次摘要：批次ID、分录笔数、计算时间')
add_para('• 展开详情：逐笔分录的完整计算过程（卖方/买方/毛利率/未售比例/未实现利润）')
add_para('• 哈希追溯：每笔分录附带SHA256哈希值，形成哈希链，任何篡改可检测')

doc.add_heading('3.10 结果导出（SRS-F10）', level=2)
add_para('支持两种导出格式：')
add_para('• Excel导出：包含汇总表（KPI指标）和抵消分录明细表（分录编号/交易类型/科目/金额/备注）')
add_para('• PDF导出：格式化抵消分录报告，含异常检测结果和计算说明')

# ============ 4. 非功能需求 ============
doc.add_heading('4. 非功能性需求', level=1)

doc.add_heading('4.1 性能要求', level=2)
add_table(
    ['指标', '要求'],
    [
        ['Excel导入', '≤ 5秒（1000行以内）'],
        ['抵消计算', '≤ 3秒（100笔交易）'],
        ['结果导出', '≤ 5秒（Excel + PDF）'],
    ]
)

doc.add_heading('4.2 可运行性', level=2)
add_para('• 单机运行，SQLite内嵌，无需数据库安装')
add_para('• pip install -r requirements.txt 一键安装依赖')
add_para('• streamlit run app.py 一键启动')
add_para('• 支持Windows/macOS/Linux')

doc.add_heading('4.3 技术栈', level=2)
add_table(
    ['层次', '选型', '说明'],
    [
        ['前端', 'Streamlit 1.x', '纯Python Web界面'],
        ['计算引擎', 'Python + JSON规则', '可配置规则引擎'],
        ['数据处理', 'pandas + openpyxl', 'Excel解析与DataFrame操作'],
        ['数据库', 'SQLite', '零配置本地存储'],
        ['图表', 'Matplotlib', '柱状图可视化'],
        ['PDF生成', 'ReportLab', '中文PDF报告'],
        ['版本控制', 'Git', '代码版本管理'],
    ]
)

# ============ 5. 数据模型 ============
doc.add_heading('5. 数据模型', level=1)

doc.add_heading('5.1 数据库表设计', level=2)
add_para('equity_config（股权关系配置）：parent_code, child_code, share_ratio, is_consolidated')
add_para('calc_log（计算日志）：elim_no, batch_id, trade_json, calc_json, entries_json, log_hash, created_at')

doc.add_heading('5.2 JSON规则配置（rules.json）', level=2)
add_para('entry_templates.DOWN/UP/FLAT 三套模板，每套含 entries（核心分录）和 minority_entries（少数股东分录）')
add_para('每条分录规则包含：dr（借方科目）、cr（贷方科目）、amount_formula（计算公式）、note（备注）')
add_para('计算公式使用安全命名空间求值，受限变量集（revenue/cost/unrealized_profit/unsold_ratio/tax_rate/minority_ratio），防止代码注入。')

# ============ 6. 接口设计 ============
doc.add_heading('6. 接口设计', level=1)
add_para('V2.0为原型验证阶段，不涉及外部系统API对接。所有数据通过Excel文件导入。')
add_para('以下为未来扩展预留接口：')
add_table(
    ['接口', '对接系统', '用途'],
    [
        ['RESTful API', 'T10智慧核算引擎', '自动获取账务数据（替代Excel上传）'],
        ['RESTful API', 'T05主数据管理', '自动获取股权关系（替代配置表）'],
        ['区块链存证API', 'T33区块链服务', '哈希上链（替代本地哈希链）'],
        ['OpenAPI 3.0', 'T21会计穿透监管', '抵消数据查询（替代文件导出）'],
    ]
)

# ============ 附录 ============
doc.add_heading('附录A：完整计算示例', level=1)
add_para('场景：母公司P（持股子公司S 80%）向S销售A产品100件，售价100万元，成本60万元。S期末对外售出30件，剩余70件。所得税税率25%。')
add_para('步骤1：交易类型识别 → P是S的母公司 → 顺流交易（DOWN）')
add_para('步骤2：未实现利润计算')
add_para('  毛利率 = (1,000,000 - 600,000) / 1,000,000 = 40%')
add_para('  未售比例 = 70 / 100 = 70%')
add_para('  未实现利润 = 1,000,000 × 40% × 70% = 280,000 元')
add_para('步骤3：抵消分录')
add_para('  借：营业收入 700,000    贷：营业成本 420,000    贷：存货 280,000')
add_para('  借：递延所得税资产 70,000    贷：所得税费用 70,000')
add_para('  合并净利润影响：-210,000 元（归母）')

doc.add_heading('附录B：V4.0核心改动说明', level=1)
add_para('V4.0在V3.0基础上新增以下功能：')
add_para('1. 试算平衡表：自动汇总借贷差额验证')
add_para('2. 分录筛选：按交易类型过滤分录卡片')
add_para('3. 分录手动调整：修改金额/删除/追加')
add_para('4. 驾驶舱增强：快捷操作卡片+最近操作时间线')
add_para('5. 历史记录：按批次查看历次计算完整过程+哈希追溯')

# ── 保存 ──
output = Path(__file__).parent / 'T11-需求说明书-V4.0.docx'
doc.save(str(output))
print(f'✅ 需求说明书已生成：{output}')
